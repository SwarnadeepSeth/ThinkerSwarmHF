import argparse
import base64
import json
import os
import re
import time
from dotenv import load_dotenv
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from agents.print_utils import render_kv_table
from tools.financial_db import set_financial_db_path, resolve_financial_db_path
from tools.sentiment_tools import set_sentiment_db_path, resolve_sentiment_db_path

try:
    from plotter.snapshot import take_snapshot
    SNAPSHOT_AVAILABLE = True
except Exception:
    SNAPSHOT_AVAILABLE = False

# Load environment variables (NVIDIA_API_KEY)
load_dotenv()

# Import the compiled LangGraph application
from core.graph import trading_app


def print_header(title: str, verbose: bool = False):
    width = max(60, len(title) + 8)
    print()
    print("╔" + "═" * (width - 2) + "╗")
    print(f"║ {title:<{width - 4}} ║")
    print("╚" + "═" * (width - 2) + "╝")
    if verbose:
        print("[VERBOSE MODE ENABLED]")


def _clean_text(value) -> str:
    return re.sub(r"\s+", " ", str(value)).strip()


def _extract_report_sections(report_text: str, section_names: list[str], max_chars: int = 140):
    sections = []
    for section_name in section_names:
        pattern = re.compile(
            rf"##\s+{re.escape(section_name)}[^\n]*\n(.*?)(?=\n##|\Z)",
            re.IGNORECASE | re.DOTALL,
        )
        match = pattern.search(report_text or "")
        if not match:
            continue
        text = _clean_text(match.group(1))
        sections.append((section_name, text[:max_chars]))
    return sections


def _extract_tool_rows(tool_results):
    rows = []
    if isinstance(tool_results, dict):
        for tool_name, result in tool_results.items():
            value = result if isinstance(result, str) else json.dumps(result, ensure_ascii=False)
            value = _clean_text(value)
            if len(value) > 220:
                value = value[:220] + "…"
            rows.append([tool_name, value])
    return rows


def _collect_tool_summary_rows(node_outputs: dict) -> list[list[str]]:
    rows = []
    if not isinstance(node_outputs, dict):
        return rows
    for node_name, payload in node_outputs.items():
        tool_results = payload.get("tool_results", {}) if isinstance(payload, dict) else {}
        if not isinstance(tool_results, dict) or not tool_results:
            continue
        tool_names = ", ".join(tool_results.keys())
        rows.append([node_name, tool_names, str(len(tool_results))])
    return rows


def _split_lines(text: str) -> list[str]:
    return [line.rstrip() for line in str(text or "").splitlines()]


def _add_verbatim_section(doc, title: str, text: str):
    doc.add_heading(title, level=1)
    content = str(text or "").strip()
    if not content:
        doc.add_paragraph("No data available.")
        return
    for line in _split_lines(content):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.font.size = Pt(8)
        run.font.name = "Courier New"


def _md_escape(value) -> str:
    text = _clean_text(value)
    return text.replace("|", "\\|").replace("\n", " ")


def _md_table(headers: list[str], rows: list[list[str]]) -> str:
    safe_headers = [_md_escape(h) for h in headers]
    lines = [
        "| " + " | ".join(safe_headers) + " |",
        "| " + " | ".join(["---"] * len(safe_headers)) + " |",
    ]
    for row in rows:
        safe_row = [_md_escape(cell) for cell in row]
        if len(safe_row) < len(safe_headers):
            safe_row.extend([""] * (len(safe_headers) - len(safe_row)))
        lines.append("| " + " | ".join(safe_row[: len(safe_headers)]) + " |")
    return "\n".join(lines)


def _truncate_for_md(value, max_chars: int = 220) -> str:
    text = _clean_text(value)
    if len(text) > max_chars:
        return text[:max_chars] + "…"
    return text


def _set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_text(cell, text, bold: bool = False, size: int = 9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)


def _style_table(table, widths=None):
    table.style = "Table Grid"
    table.autofit = False
    if widths:
        for idx, width in enumerate(widths):
            table.columns[idx].width = Inches(width)


def add_table_section(doc, title: str, headers: list[str], rows: list[list[str]], widths=None):
    doc.add_heading(title, level=1)
    if not rows:
        doc.add_paragraph("No data available.")
        return

    table = doc.add_table(rows=1, cols=len(headers))
    _style_table(table, widths)
    header_row = table.rows[0].cells
    for idx, header in enumerate(headers):
        _set_cell_text(header_row[idx], header, bold=True, size=10)
        _set_cell_shading(header_row[idx], "1F4E79")
        for paragraph in header_row[idx].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    for row in rows:
        cells = table.add_row().cells
        padded = list(row) + [""] * max(0, len(headers) - len(row))
        for idx, value in enumerate(padded[: len(headers)]):
            _set_cell_text(cells[idx], value)

    doc.add_paragraph()


def generate_docx_report(ticker: str, out_data: dict, decision: dict):
    """Generate a DOCX report from the analysis results."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Title
    title = doc.add_heading(f"Trading Analysis Report: {ticker}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_rows = [
        ["Generated", time.strftime("%Y-%m-%d %H:%M:%S")],
        ["Execution Time", f"{out_data.get('execution_time_ms', 0):.2f} ms"],
        ["Ticker", ticker],
        ["Nodes Executed", ", ".join(out_data.get("nodes_executed", [])) or "None"],
        ["Status", "ERROR" if out_data.get("error") else "SUCCESS"],
    ]
    if out_data.get("error"):
        meta_rows.append(["Error", out_data.get("error")])
    add_table_section(doc, "Run Metadata", ["Field", "Value"], meta_rows, widths=[1.7, 4.9])

    # Executive Summary / Final Decision
    if decision:
        entry = float(decision.get("entry_price", 0) or 0)
        stop = float(decision.get("stoploss", 0) or 0)
        target = float(decision.get("target", 0) or 0)
        rr = ""
        if entry and stop and target:
            risk = abs(entry - stop)
            reward = abs(target - entry)
            if risk:
                rr = f"{reward / risk:.2f}x"

        decision_rows = [
            ["Direction", decision.get("direction", "N/A")],
            ["Entry", f"${entry:,.2f}" if entry else "N/A"],
            ["Stop", f"${stop:,.2f}" if stop else "N/A"],
            ["Target", f"${target:,.2f}" if target else "N/A"],
            ["Risk / Reward", rr or "N/A"],
            ["Volatility", decision.get("risk_volatility", "N/A")],
            ["Timeframe", decision.get("timeframe", "N/A")],
            ["Reasoning", decision.get("reasoning", "")],
        ]
        add_table_section(doc, "Final Trade Setup", ["Field", "Value"], decision_rows, widths=[1.7, 4.9])
    else:
        doc.add_paragraph("No decision generated.")

    tech_report = ""
    fund_report = ""
    for step in out_data.get("steps", []):
        if step.get("node") == "quant_head_synthesis":
            tech_report = step.get("llm_output", "") or ""
        elif step.get("node") == "fund_head_synthesis":
            fund_report = step.get("llm_output", "") or ""

    if tech_report:
        add_table_section(
            doc,
            "Technical Wing Highlights",
            ["Section", "Excerpt"],
            _extract_report_sections(
                tech_report,
                [
                    "Rationale",
                    "Bull Perspective",
                    "Bear Perspective",
                    "Risks",
                    "Opportunities",
                    "Priority",
                    "Recommended Stop Loss",
                    "Profit Target",
                    "Time Horizon",
                    "Overall Bias",
                ],
            ),
            widths=[2.2, 4.4],
        )

    if fund_report:
        add_table_section(
            doc,
            "Fundamental Wing Highlights",
            ["Section", "Excerpt"],
            _extract_report_sections(
                fund_report,
                [
                    "Rationale",
                    "Bull Perspective",
                    "Bear Perspective",
                    "Risks",
                    "Opportunities",
                    "Priority",
                    "Recommended Stop Loss",
                    "Profit Target",
                    "Time Horizon",
                    "Overall Bias",
                ],
            ),
            widths=[2.2, 4.4],
        )

    fundamental_analysis = out_data.get("fundamental_analysis", {}) or {}
    sector_peer_context = fundamental_analysis.get("sector_peer_context", "")
    internet_context = fundamental_analysis.get("internet_context", "")
    if sector_peer_context:
        _add_verbatim_section(doc, "Sector & Peer Comparative Context", sector_peer_context)
    if internet_context:
        _add_verbatim_section(doc, "Internet Search Context", internet_context)
    sentiment_analysis = out_data.get("sentiment_analysis", {}) or {}
    raw_sentiment_context = sentiment_analysis.get("raw_sentiment_context", "")
    sentiment_brief = sentiment_analysis.get("sentiment_context", "")
    if raw_sentiment_context:
        _add_verbatim_section(doc, "Sentiment Database + Live News Context", raw_sentiment_context)
    if sentiment_brief:
        _add_verbatim_section(doc, "Sentiment Wing Brief", sentiment_brief)

    node_outputs = out_data.get("node_outputs", {}) or {}
    tool_sections = [
        ("Quantitative Bull Tool Stack", node_outputs.get("quant_bull_worker", {}).get("tool_results", {})),
        ("Quantitative Bear Tool Stack", node_outputs.get("quant_bear_worker", {}).get("tool_results", {})),
        ("Fundamental Bull Tool Stack", node_outputs.get("fund_bull_worker", {}).get("tool_results", {})),
        ("Fundamental Bear Tool Stack", node_outputs.get("fund_bear_worker", {}).get("tool_results", {})),
        ("Quant Core Tool Stack", node_outputs.get("quant", {}).get("tool_results", {})),
    ]

    for title, tool_results in tool_sections:
        rows = _extract_tool_rows(tool_results)
        if rows:
            add_table_section(doc, title, ["Tool", "Result"], rows, widths=[2.4, 4.2])

    # Analysis Steps
    doc.add_heading("Analysis Steps", level=1)
    step_rows = []
    for step in out_data.get("steps", []):
        output = step.get("llm_output", "")
        if isinstance(output, dict):
            output = json.dumps(output, ensure_ascii=False)
        output = _clean_text(output)
        if len(output) > 180:
            output = output[:180] + "…"
        step_rows.append(
            [
                step.get("node", "unknown"),
                step.get("model_used", "unknown"),
                f"{step.get('timestamp', 0):.2f}s",
                output,
            ]
        )
    add_table_section(
        doc,
        "Execution Steps",
        ["Node", "Model", "Duration", "Preview"],
        step_rows,
        widths=[1.25, 1.35, 0.75, 3.75],
    )

    # Append full wing reports verbatim for auditability
    doc.add_heading("Appendix - Full Technical Wing Report", level=1)
    if tech_report:
        doc.add_paragraph(tech_report)
    else:
        doc.add_paragraph("No technical wing report available.")

    doc.add_heading("Appendix - Full Fundamental Wing Report", level=1)
    if fund_report:
        doc.add_paragraph(fund_report)
    else:
        doc.add_paragraph("No fundamental wing report available.")

    tool_summary_rows = _collect_tool_summary_rows(node_outputs)
    if tool_summary_rows:
        add_table_section(
            doc,
            "Appendix - Tool Calls by Node",
            ["Node", "Tools Used", "Count"],
            tool_summary_rows,
            widths=[2.0, 4.0, 0.8],
        )

    # Save
    filename = f"{ticker}_analysis_report.docx"
    doc.save(filename)
    print(f"📄 DOCX report saved to {filename}")


def generate_markdown_report(ticker: str, out_data: dict, decision: dict):
    """Generate a Markdown report from the analysis results."""
    lines = [f"# Trading Analysis Report: {ticker}", ""]

    meta_rows = [
        ["Generated", time.strftime("%Y-%m-%d %H:%M:%S")],
        ["Execution Time", f"{out_data.get('execution_time_ms', 0):.2f} ms"],
        ["Ticker", ticker],
        ["Nodes Executed", ", ".join(out_data.get("nodes_executed", [])) or "None"],
        ["Status", "ERROR" if out_data.get("error") else "SUCCESS"],
    ]
    if out_data.get("error"):
        meta_rows.append(["Error", out_data.get("error")])
    lines.extend(["## Run Metadata", _md_table(["Field", "Value"], meta_rows), ""])

    if decision:
        entry = float(decision.get("entry_price", 0) or 0)
        stop = float(decision.get("stoploss", 0) or 0)
        target = float(decision.get("target", 0) or 0)
        rr = ""
        if entry and stop and target:
            risk = abs(entry - stop)
            reward = abs(target - entry)
            if risk:
                rr = f"{reward / risk:.2f}x"

        decision_rows = [
            ["Direction", decision.get("direction", "N/A")],
            ["Entry", f"${entry:,.2f}" if entry else "N/A"],
            ["Stop", f"${stop:,.2f}" if stop else "N/A"],
            ["Target", f"${target:,.2f}" if target else "N/A"],
            ["Risk / Reward", rr or "N/A"],
            ["Volatility", decision.get("risk_volatility", "N/A")],
            ["Timeframe", decision.get("timeframe", "N/A")],
            ["Reasoning", decision.get("reasoning", "")],
        ]
        lines.extend(["## Final Trade Setup", _md_table(["Field", "Value"], decision_rows), ""])
    else:
        lines.extend(["## Final Trade Setup", "No decision generated.", ""])

    tech_report = ""
    fund_report = ""
    for step in out_data.get("steps", []):
        if step.get("node") == "quant_head_synthesis":
            tech_report = step.get("llm_output", "") or ""
        elif step.get("node") == "fund_head_synthesis":
            fund_report = step.get("llm_output", "") or ""

    if tech_report:
        tech_rows = _extract_report_sections(
            tech_report,
            [
                "Rationale",
                "Bull Perspective",
                "Bear Perspective",
                "Risks",
                "Opportunities",
                "Priority",
                "Recommended Stop Loss",
                "Profit Target",
                "Time Horizon",
                "Overall Bias",
            ],
        )
        if tech_rows:
            tech_rows = [[sec, _truncate_for_md(text, 300)] for sec, text in tech_rows]
            lines.extend(["## Technical Wing Highlights", _md_table(["Section", "Excerpt"], tech_rows), ""])

    if fund_report:
        fund_rows = _extract_report_sections(
            fund_report,
            [
                "Rationale",
                "Bull Perspective",
                "Bear Perspective",
                "Risks",
                "Opportunities",
                "Priority",
                "Recommended Stop Loss",
                "Profit Target",
                "Time Horizon",
                "Overall Bias",
            ],
        )
        if fund_rows:
            fund_rows = [[sec, _truncate_for_md(text, 300)] for sec, text in fund_rows]
            lines.extend(["## Fundamental Wing Highlights", _md_table(["Section", "Excerpt"], fund_rows), ""])

    fundamental_analysis = out_data.get("fundamental_analysis", {}) or {}
    sector_peer_context = fundamental_analysis.get("sector_peer_context", "")
    internet_context = fundamental_analysis.get("internet_context", "")
    if sector_peer_context:
        lines.extend(["## Sector & Peer Comparative Context", "```"])
        lines.append(sector_peer_context)
        lines.extend(["```", ""])
    if internet_context:
        lines.extend(["## Internet Search Context", "```"])
        lines.append(internet_context)
        lines.extend(["```", ""])
    sentiment_analysis = out_data.get("sentiment_analysis", {}) or {}
    raw_sentiment_context = sentiment_analysis.get("raw_sentiment_context", "")
    sentiment_brief = sentiment_analysis.get("sentiment_context", "")
    if raw_sentiment_context:
        lines.extend(["## Sentiment Database + Live News Context", "```"])
        lines.append(raw_sentiment_context)
        lines.extend(["```", ""])
    if sentiment_brief:
        lines.extend(["## Sentiment Wing Brief", "```"])
        lines.append(sentiment_brief)
        lines.extend(["```", ""])

    node_outputs = out_data.get("node_outputs", {}) or {}
    tool_sections = [
        ("Quantitative Bull Tool Stack", node_outputs.get("quant_bull_worker", {}).get("tool_results", {})),
        ("Quantitative Bear Tool Stack", node_outputs.get("quant_bear_worker", {}).get("tool_results", {})),
        ("Fundamental Bull Tool Stack", node_outputs.get("fund_bull_worker", {}).get("tool_results", {})),
        ("Fundamental Bear Tool Stack", node_outputs.get("fund_bear_worker", {}).get("tool_results", {})),
        ("Quant Core Tool Stack", node_outputs.get("quant", {}).get("tool_results", {})),
    ]
    for section_title, tool_results in tool_sections:
        rows = _extract_tool_rows(tool_results)
        if rows:
            lines.extend([f"## {section_title}", _md_table(["Tool", "Result"], rows), ""])

    step_rows = []
    for step in out_data.get("steps", []):
        output = step.get("llm_output", "")
        if isinstance(output, dict):
            output = json.dumps(output, ensure_ascii=False)
        step_rows.append(
            [
                step.get("node", "unknown"),
                step.get("model_used", "unknown"),
                f"{step.get('timestamp', 0):.2f}s",
                _truncate_for_md(output, 220),
            ]
        )
    lines.extend(["## Execution Steps", _md_table(["Node", "Model", "Duration", "Preview"], step_rows), ""])

    lines.extend(["## Appendix - Full Technical Wing Report", "```"])
    lines.append(tech_report or "No technical wing report available.")
    lines.extend(["```", ""])

    lines.extend(["## Appendix - Full Fundamental Wing Report", "```"])
    lines.append(fund_report or "No fundamental wing report available.")
    lines.extend(["```", ""])

    tool_summary_rows = _collect_tool_summary_rows(node_outputs)
    if tool_summary_rows:
        lines.extend([
            "## Appendix - Tool Calls by Node",
            _md_table(["Node", "Tools Used", "Count"], tool_summary_rows),
            "",
        ])

    filename = f"{ticker}_analysis_report.md"
    with open(filename, "w", encoding="utf-8") as f:
        f.write("\n".join(lines).rstrip() + "\n")
    print(f"📝 Markdown report saved to {filename}")


def generate_html_report(ticker: str, out_data: dict, decision: dict):
    """Generate a styled HTML report with chart snapshot - includes all sections from MD/DOCX."""
    chart_img = ""
    if SNAPSHOT_AVAILABLE:
        try:
            db_path = "data/US_DB.db"
            png_path = take_snapshot(ticker, db_path)
            if png_path and os.path.exists(png_path):
                with open(png_path, "rb") as f:
                    chart_img = base64.b64encode(f.read()).decode("utf-8")
        except Exception as e:
            print(f"⚠️ Chart snapshot failed: {e}")

    def row_html(label: str, value: str) -> str:
        return f'<tr><td class="label">{label}</td><td class="value">{value}</td></tr>'

    def escape_html(text: str) -> str:
        if not text:
            return ""
        return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("|", "│")

    meta_rows = row_html("Generated", time.strftime("%Y-%m-%d %H:%M:%S"))
    meta_rows += row_html("Execution Time", f"{out_data.get('execution_time_ms', 0):.2f} ms")
    meta_rows += row_html("Ticker", ticker)
    meta_rows += row_html("Nodes Executed", ", ".join(out_data.get("nodes_executed", [])) or "None")
    meta_rows += row_html("Status", f'<span class="status-{"error" if out_data.get("error") else "success"}">{"ERROR" if out_data.get("error") else "SUCCESS"}</span>')
    if out_data.get("error"):
        meta_rows += row_html("Error", out_data.get("error"))

    decision_html = ""
    if decision:
        entry = float(decision.get("entry_price", 0) or 0)
        stop = float(decision.get("stoploss", 0) or 0)
        target = float(decision.get("target", 0) or 0)
        rr = ""
        if entry and stop and target:
            risk = abs(entry - stop)
            reward = abs(target - entry)
            if risk:
                rr = f"{reward / risk:.2f}x"
        reasoning_text = decision.get("reasoning", "")
        decision_rows = [
            ("Direction", decision.get("direction", "N/A")),
            ("Entry", f"${entry:,.2f}" if entry else "N/A"),
            ("Stop", f"${stop:,.2f}" if stop else "N/A"),
            ("Target", f"${target:,.2f}" if target else "N/A"),
            ("Risk / Reward", rr or "N/A"),
            ("Volatility", decision.get("risk_volatility", "N/A")),
            ("Timeframe", decision.get("timeframe", "N/A")),
            ("Reasoning", f'<div class="reasoning-box">{escape_html(reasoning_text)}</div>'),
        ]
        decision_html = "\n".join(row_html(label, val) for label, val in decision_rows)
    else:
        decision_html = '<tr><td colspan="2">No decision generated.</td></tr>'

    tech_report = ""
    fund_report = ""
    for step in out_data.get("steps", []):
        if step.get("node") == "quant_head_synthesis":
            tech_report = step.get("llm_output", "") or ""
        elif step.get("node") == "fund_head_synthesis":
            fund_report = step.get("llm_output", "") or ""

    tech_sections = []
    for sec, txt in _extract_report_sections(tech_report, ["Rationale", "Bull Perspective", "Bear Perspective", "Risks", "Opportunities", "Priority", "Recommended Stop Loss", "Profit Target", "Time Horizon", "Overall Bias"]):
        tech_sections.append(f"<tr><td class='section'>{sec}</td><td>{escape_html(txt)}</td></tr>")
    tech_html = "\n".join(tech_sections) if tech_sections else "<tr><td colspan='2'>No technical analysis available.</td></tr>"

    fund_sections = []
    for sec, txt in _extract_report_sections(fund_report, ["Rationale", "Bull Perspective", "Bear Perspective", "Risks", "Opportunities", "Priority", "Recommended Stop Loss", "Profit Target", "Time Horizon", "Overall Bias"]):
        fund_sections.append(f"<tr><td class='section'>{sec}</td><td>{escape_html(txt)}</td></tr>")
    fund_html = "\n".join(fund_sections) if fund_sections else "<tr><td colspan='2'>No fundamental analysis available.</td></tr>"

    fundamental_analysis = out_data.get("fundamental_analysis", {}) or {}
    sector_peer_context = fundamental_analysis.get("sector_peer_context", "")
    internet_context = fundamental_analysis.get("internet_context", "")

    sector_html = ""
    if sector_peer_context:
        sector_html = f'<div class="code-block"><pre>{escape_html(sector_peer_context)}</pre></div>'
    else:
        sector_html = '<p class="empty">No sector/peer context available.</p>'

    internet_html = ""
    if internet_context:
        internet_html = f'<div class="code-block"><pre>{escape_html(internet_context)}</pre></div>'
    else:
        internet_html = '<p class="empty">No internet search context available.</p>'

    sentiment_analysis = out_data.get("sentiment_analysis", {}) or {}
    raw_sentiment_context = sentiment_analysis.get("raw_sentiment_context", "")
    sentiment_brief = sentiment_analysis.get("sentiment_context", "")

    sentiment_html = ""
    if raw_sentiment_context:
        sentiment_html = f'<div class="code-block"><pre>{escape_html(raw_sentiment_context)}</pre></div>'
    else:
        sentiment_html = '<p class="empty">No sentiment data available.</p>'

    sentiment_brief_html = ""
    if sentiment_brief:
        sentiment_brief_html = f'<div class="code-block"><pre>{escape_html(sentiment_brief)}</pre></div>'
    else:
        sentiment_brief_html = '<p class="empty">No sentiment brief available.</p>'

    node_outputs = out_data.get("node_outputs", {}) or {}
    tool_html = ""
    tool_sections = [
        ("Quantitative Bull Tool Stack", node_outputs.get("quant_bull_worker", {}).get("tool_results", {})),
        ("Quantitative Bear Tool Stack", node_outputs.get("quant_bear_worker", {}).get("tool_results", {})),
        ("Fundamental Bull Tool Stack", node_outputs.get("fund_bull_worker", {}).get("tool_results", {})),
        ("Fundamental Bear Tool Stack", node_outputs.get("fund_bear_worker", {}).get("tool_results", {})),
        ("Quant Core Tool Stack", node_outputs.get("quant", {}).get("tool_results", {})),
    ]
    for name, results in tool_sections:
        rows = _extract_tool_rows(results)
        if rows:
            tool_html += f"<h3>{name}</h3>\n<table class='tools'><thead><tr><th>Tool</th><th>Result</th></tr></thead><tbody>\n"
            for tool, result in rows:
                tool_html += f"<tr><td class='tool'>{escape_html(tool)}</td><td>{escape_html(result)}</td></tr>\n"
            tool_html += "</tbody></table>\n"

    step_rows = []
    for step in out_data.get("steps", []):
        output = step.get("llm_output", "")
        if isinstance(output, dict):
            output = json.dumps(output, ensure_ascii=False)
        step_rows.append([
            step.get("node", "unknown"),
            step.get("model_used", "unknown"),
            f"{step.get('timestamp', 0):.2f}s",
            _truncate_for_md(output, 220) if output else "",
        ])
    steps_html = ""
    if step_rows:
        steps_html = "<table class='steps'><thead><tr><th>Node</th><th>Model</th><th>Duration</th><th>Preview</th></tr></thead><tbody>\n"
        for node, model, duration, preview in step_rows:
            steps_html += f"<tr><td>{escape_html(node)}</td><td>{escape_html(model)}</td><td>{escape_html(duration)}</td><td>{escape_html(preview)}</td></tr>\n"
        steps_html += "</tbody></table>"

    tech_full_html = f'<div class="code-block full-report"><pre>{escape_html(tech_report)}</pre></div>' if tech_report else '<p class="empty">No technical wing report available.</p>'
    fund_full_html = f'<div class="code-block full-report"><pre>{escape_html(fund_report)}</pre></div>' if fund_report else '<p class="empty">No fundamental wing report available.</p>'

    tool_summary_rows = _collect_tool_summary_rows(node_outputs)
    tool_summary_html = ""
    if tool_summary_rows:
        tool_summary_html = "<table class='tool-summary'><thead><tr><th>Node</th><th>Tools Used</th><th>Count</th></tr></thead><tbody>\n"
        for node_name, tool_names, count in tool_summary_rows:
            tool_summary_html += f"<tr><td>{escape_html(node_name)}</td><td>{escape_html(tool_names)}</td><td>{escape_html(count)}</td></tr>\n"
        tool_summary_html += "</tbody></table>"

    html = f'''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Trading Analysis Report: {ticker}</title>
<style>
    * {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif; background: #0b0f1a; color: #e6edf3; line-height: 1.6; padding: 40px; }}
    .container {{ max-width: 1200px; margin: 0 auto; }}
    h1 {{ color: #58a6ff; font-size: 28px; margin-bottom: 8px; }}
    h2 {{ color: #f5d300; font-size: 18px; margin: 30px 0 15px; border-bottom: 1px solid #30363d; padding-bottom: 8px; }}
    h3 {{ color: #8b949e; font-size: 14px; margin: 20px 0 10px; text-transform: uppercase; letter-spacing: 1px; }}
    table {{ width: 100%; border-collapse: collapse; margin: 15px 0; background: #161b22; border-radius: 8px; overflow: hidden; }}
    th, td {{ padding: 12px 16px; text-align: left; border-bottom: 1px solid #30363d; }}
    th {{ background: #21262d; color: #8b949e; font-size: 11px; text-transform: uppercase; letter-spacing: 1px; }}
    td.label {{ color: #8b949e; width: 180px; font-weight: 500; }}
    td.value {{ color: #e6edf3; }}
    td.section {{ color: #58a6ff; width: 160px; }}
    td.tool {{ color: #c9a0fa; width: 140px; }}
    .status-success {{ color: #3fb950; font-weight: bold; }}
    .status-error {{ color: #f85149; font-weight: bold; }}
    .chart-container {{ margin: 30px 0; text-align: center; }}
    .chart-container img {{ max-width: 100%; border-radius: 12px; box-shadow: 0 8px 32px rgba(0,0,0,0.4); }}
    .meta {{ color: #8b949e; font-size: 13px; margin-bottom: 30px; }}
    .reasoning-box {{ background: #161b22; padding: 16px; border-radius: 8px; border-left: 3px solid #f5d300; margin-top: 8px; white-space: pre-wrap; }}
    .code-block {{ background: #161b22; padding: 16px; border-radius: 8px; overflow-x: auto; margin: 15px 0; }}
    .code-block pre {{ margin: 0; font-size: 12px; font-family: "SF Mono", Monaco, "Courier New", monospace; color: #8b949e; white-space: pre-wrap; word-wrap: break-word; }}
    .full-report pre {{ color: #e6edf3; max-height: 400px; overflow-y: auto; }}
    .empty {{ color: #8b949e; font-style: italic; }}
    .footer {{ margin-top: 40px; padding-top: 20px; border-top: 1px solid #30363d; color: #8b949e; font-size: 12px; text-align: center; }}
    .appendix {{ margin-top: 40px; border-top: 2px dashed #30363d; padding-top: 20px; }}
    .appendix h2 {{ color: #8b949e; }}
    @media print {{ body {{ background: white; color: black; }} table {{ background: white; }} .code-block {{ background: #f5f5f5; }} }}
</style>
</head>
<body>
<div class="container">
    <h1>⚡ Trading Analysis Report: {ticker}</h1>
    <p class="meta">Generated {time.strftime("%Y-%m-%d at %H:%M:%S")}</p>
'''

    if chart_img:
        html += f'    <div class="chart-container"><img src="data:image/png;base64,{chart_img}" alt="Price Chart"></div>\n'

    html += f'''    <h2>Run Metadata</h2>
    <table><tbody>{meta_rows}</tbody></table>

    <h2>Final Trade Setup</h2>
    <table><tbody>{decision_html}</tbody></table>

    <h2>Technical Wing Highlights</h2>
    <table><tbody>{tech_html}</tbody></table>

    <h2>Fundamental Wing Highlights</h2>
    <table><tbody>{fund_html}</tbody></table>

    <h2>Sector & Peer Comparative Context</h2>
    {sector_html}

    <h2>Internet Search Context</h2>
    {internet_html}

    <h2>Sentiment Database + Live News Context</h2>
    {sentiment_html}

    <h2>Sentiment Wing Brief</h2>
    {sentiment_brief_html}

    <h2>Tool Results</h2>
'''
    html += tool_html

    html += f'''    <h2>Execution Steps</h2>
    {steps_html}

    <div class="appendix">
    <h2>Appendix - Full Technical Wing Report</h2>
    {tech_full_html}

    <h2>Appendix - Full Fundamental Wing Report</h2>
    {fund_full_html}
'''

    if tool_summary_html:
        html += f'''    <h2>Appendix - Tool Calls by Node</h2>
    {tool_summary_html}
'''

    html += f'''    <div class="footer">Generated by ThinkerSwarm HF • QuantJuice Pro</div>
</div>
</body>
</html>'''

    filename = f"{ticker}_analysis_report.html"
    with open(filename, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"📊 HTML report saved to {filename}")


def main():
    parser = argparse.ArgumentParser(
        description="Autonomous AI Trading Agent Framework"
    )
    parser.add_argument(
        "--ticker", type=str, required=True, help="Stock ticker to analyze (e.g., AAPL)"
    )
    parser.add_argument(
        "--db", type=str, default="data/US_DB.db", help="Path to your SQLite database"
    )
    parser.add_argument(
        "--financial-db",
        type=str,
        default="",
        help="Path to the fundamental SQLite database (default: data/financials.db)",
    )
    parser.add_argument(
        "--sentiment-db",
        type=str,
        default="",
        help="Path to the sentiment SQLite database (default: data/sentiments.db)",
    )
    parser.add_argument(
        "-v", "--verbose", action="store_true", help="Enable verbose debug output"
    )
    parser.add_argument(
        "--research-web",
        action="store_true",
        help="Enable free internet search context in the Researcher node",
    )

    args = parser.parse_args()

    # 1. Validate the database exists
    if not os.path.exists(args.db):
        print(f"❌ Error: Database file not found at '{args.db}'")
        print("Please ensure your 'US_DB.db' is placed inside the 'data/' folder.")
        return

    financial_db_path = resolve_financial_db_path(args.financial_db or None)
    if not financial_db_path:
        print("⚠️ Warning: financial database not found. Fundamental tools will fall back to yfinance.")
        financial_db_path = args.financial_db or "data/financials.db"
    set_financial_db_path(financial_db_path)

    sentiment_db_path = resolve_sentiment_db_path(args.sentiment_db or None)
    if not sentiment_db_path:
        print("⚠️ Warning: sentiment database not found. Sentiment wing will fall back to live news only.")
        sentiment_db_path = args.sentiment_db or "data/sentiments.db"
    set_sentiment_db_path(sentiment_db_path)

    print_header(f"🚀 Initializing Framework for {args.ticker.upper()}", args.verbose)
    print(f"🔌 Connected to database: {args.db}")
    print(f"💼 Financial database: {financial_db_path}")
    print(f"🧠 Sentiment database: {sentiment_db_path}")
    print(f"🌐 Research web search: {'ENABLED' if args.research_web else 'DISABLED'}")
    if args.verbose:
        print(f"📋 Verbose mode: ENABLED")
    print(
        "🧠 Agents are starting their analysis... (This may take a few minutes if new indicators are being coded)"
    )

    # 2. Initialize the Shared State
    execution_start = time.time()
    initial_state = {
        # Base
        "ticker": args.ticker.upper(),
        "db_path": args.db,
        "financial_db_path": financial_db_path,
        "sentiment_db_path": sentiment_db_path,
        "verbose": args.verbose,
        "allow_research_web": args.research_web,
        "iteration_count": 0,
        "market_context": "",
        "manager_brief": "",
        # Researcher
        "researcher_context": "",
        # Sentiment wing
        "sentiment_context": "",
        "sentiment_analysis": {},
        # Quantitative wing
        "quant_bull_findings": "",
        "quant_bear_findings": "",
        "quant_wing_report": "",
        # Fundamental wing
        "fund_bull_findings": "",
        "fund_bear_findings": "",
        "fund_wing_report": "",
        # Coder fallback
        "coder_loop_count": 0,
        "indicator_requested": "",
        "draft_code": "",
        "test_results": "",
        "code_review_feedback": "",
        "code_approved": False,
        "quant_used_tools": False,
        "quant_strategy": "",
        "coder_output": {},
        # Legacy analysis dicts (for reviewer compat)
        "technical_analysis": {},
        "fundamental_analysis": {},
        # Review
        "reviewer_feedback": "",
        "final_decision": {},
        # Tracking
        "execution_start_time": execution_start,
        "node_timestamps": {},
        "node_outputs": {},
    }

    # 3. Execute the Graph
    error_msg = None
    try:
        final_state = trading_app.invoke(initial_state)

        print_header("🏁 FINAL TRADE SETUP APPROVED", args.verbose)

        decision = final_state.get("final_decision", {})
        if decision:
            entry = float(decision.get("entry_price", 0) or 0)
            stop = float(decision.get("stoploss", 0) or 0)
            target = float(decision.get("target", 0) or 0)
            rr = "N/A"
            if entry and stop and target:
                risk = abs(entry - stop)
                reward = abs(target - entry)
                if risk:
                    rr = f"{reward / risk:.2f}x"

            render_kv_table(
                "Final Trade Setup",
                [
                    ("Direction", decision.get("direction", "N/A")),
                    ("Entry", f"${entry:,.2f}" if entry else "N/A"),
                    ("Stop", f"${stop:,.2f}" if stop else "N/A"),
                    ("Target", f"${target:,.2f}" if target else "N/A"),
                    ("Risk / Reward", rr),
                    ("Volatility", decision.get("risk_volatility", "N/A")),
                    ("Timeframe", decision.get("timeframe", "N/A")),
                    ("Reasoning", decision.get("reasoning", "")),
                ],
            )
        else:
            print("❌ Framework completed, but no final decision was generated.")
            print("Reviewer Feedback:", final_state.get("reviewer_feedback"))

    except Exception as e:
        print_header("⚠️ EXECUTION FAILED", args.verbose)
        print(f"An error occurred during agent routing: {str(e)}")
        error_msg = str(e)
        final_state = {}

    # 4. Save output to out.json
    execution_total_time = time.time() - execution_start
    node_timestamps = final_state.get("node_timestamps", {}) if final_state else {}
    node_outputs = final_state.get("node_outputs", {}) if final_state else {}

    steps = []
    for node_name, output_data in node_outputs.items():
        step_entry = {
            "node": node_name,
            "timestamp": node_timestamps.get(node_name, 0),
            "llm_output": output_data.get("llm_output", ""),
            "execution_result": output_data.get("execution_result", ""),
            "model_used": output_data.get("model_used", "unknown"),
        }
        steps.append(step_entry)

    decision = final_state.get("final_decision", {}) if final_state else {}

    out_data = {
        "ticker": args.ticker.upper(),
        "execution_time_ms": round(execution_total_time * 1000, 2),
        "nodes_executed": list(node_outputs.keys()),
        "steps": steps,
        "node_outputs": node_outputs,
        "fundamental_analysis": final_state.get("fundamental_analysis", {}) if final_state else {},
        "sentiment_analysis": final_state.get("sentiment_analysis", {}) if final_state else {},
        "final_decision": decision,
        "error": error_msg if error_msg else None,
    }

    with open("out.json", "w") as f:
        json.dump(out_data, f, indent=2)

    print(f"\n💾 Output saved to out.json")

    # 5. Generate DOCX report
    generate_docx_report(args.ticker.upper(), out_data, decision)
    generate_markdown_report(args.ticker.upper(), out_data, decision)
    generate_html_report(args.ticker.upper(), out_data, decision)


if __name__ == "__main__":
    main()
